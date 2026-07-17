"""
RAG 知识库管理：知识库 CRUD + 文档上传/切片/嵌入/检索/问答
"""
from __future__ import annotations

import os, json
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query, Form
from fastapi.responses import JSONResponse
from database.session import SessionLocal, get_db
from core.security import get_current_user
from core.rbac import require_role
from core.upload_security import KB_EXTENSIONS, save_validated_upload
from core.milvus_client import (
    ensure_collection, insert_chunks, search_chunks, delete_doc_chunks,
    delete_collection, collection_stats, milvus_available, MILVUS_URI,
)
from core.embedding_service import get_embedding, chunk_text, rerank_documents
from dao.model_dao import get_model
from models.user import User
from models.knowledge_base import KnowledgeBase, KBDocument
from models.ai_model import AIModel
from pydantic import BaseModel

kb_router = APIRouter(prefix="/api/kb", tags=["RAG知识库"])

# ============ Schemas ============
class KBCreateIn(BaseModel):
    name: str
    description: str = ""
    embedding_model_id: int | None = None
    rerank_model_id: int | None = None
    chunk_size: int = 500
    chunk_overlap: int = 50


class KBUpdateIn(BaseModel):
    name: str | None = None
    description: str | None = None
    embedding_model_id: int | None = None
    rerank_model_id: int | None = None
    chunk_size: int | None = None
    chunk_overlap: int | None = None


class KBQueryIn(BaseModel):
    kb_id: int
    question: str
    top_k: int = 5
    use_rerank: bool = True


# ============ 知识库 CRUD ============
@kb_router.get("/status")
def kb_status():
    """检查向量数据库后端状态"""
    return {"milvus_available": milvus_available(), "milvus_uri": MILVUS_URI}


@kb_router.get("")
def list_kb(db: SessionLocal = Depends(get_db), user: User = Depends(get_current_user)):
    kbs = db.query(KnowledgeBase).order_by(KnowledgeBase.updated_at.desc()).all()
    result = []
    for kb in kbs:
        emb_name, rerank_name = "", ""
        if kb.embedding_model_id:
            m = get_model(kb.embedding_model_id, db)
            emb_name = m.name if m else ""
        if kb.rerank_model_id:
            m = get_model(kb.rerank_model_id, db)
            rerank_name = m.name if m else ""
        result.append({
            "id": kb.id, "name": kb.name, "description": kb.description,
            "embedding_model_id": kb.embedding_model_id, "embedding_model_name": emb_name,
            "rerank_model_id": kb.rerank_model_id, "rerank_model_name": rerank_name,
            "chunk_size": kb.chunk_size, "chunk_overlap": kb.chunk_overlap,
            "doc_count": kb.doc_count, "created_at": kb.created_at.isoformat(),
            "updated_at": kb.updated_at.isoformat(),
        })
    return result


@kb_router.post("", status_code=201)
def create_kb(body: KBCreateIn, db: SessionLocal = Depends(get_db),
              user: User = Depends(require_role("ROOT", "OPS", "ADMIN"))):
    kb = KnowledgeBase(
        name=body.name, description=body.description,
        embedding_model_id=body.embedding_model_id, rerank_model_id=body.rerank_model_id,
        chunk_size=body.chunk_size, chunk_overlap=body.chunk_overlap,
    )
    db.add(kb)
    db.commit()
    db.refresh(kb)
    ensure_collection(kb.id)
    return {"id": kb.id, "name": kb.name, "description": kb.description}


@kb_router.put("/{kb_id}")
def update_kb(kb_id: int, body: KBUpdateIn, db: SessionLocal = Depends(get_db),
              user: User = Depends(require_role("ROOT", "OPS", "ADMIN"))):
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
    if not kb:
        raise HTTPException(404, "知识库不存在")
    for k, v in body.model_dump().items():
        if v is not None:
            setattr(kb, k, v)
    db.commit()
    db.refresh(kb)
    return {"id": kb.id, "name": kb.name}


@kb_router.delete("/{kb_id}")
def delete_kb(kb_id: int, db: SessionLocal = Depends(get_db),
              user: User = Depends(require_role("ROOT", "OPS", "ADMIN"))):
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
    if not kb:
        raise HTTPException(404, "知识库不存在")
    # 删除 Milvus 中的向量
    delete_collection(kb_id)
    # 删除关联文档
    db.query(KBDocument).filter(KBDocument.kb_id == kb_id).delete()
    db.delete(kb)
    db.commit()
    return {"deleted": True}


# ============ 文档管理 ============
@kb_router.get("/{kb_id}/docs")
def list_docs(kb_id: int, db: SessionLocal = Depends(get_db), user: User = Depends(get_current_user)):
    docs = db.query(KBDocument).filter(KBDocument.kb_id == kb_id).order_by(KBDocument.created_at.desc()).all()
    return [{
        "id": d.id, "kb_id": d.kb_id, "filename": d.filename, "file_type": d.file_type,
        "file_size": d.file_size, "chunk_count": d.chunk_count, "status": d.status,
        "error_msg": d.error_msg, "created_at": d.created_at.isoformat(),
    } for d in docs]


@kb_router.post("/{kb_id}/upload")
async def upload_doc(kb_id: int, file: UploadFile = File(...), db: SessionLocal = Depends(get_db),
                     user: User = Depends(require_role("ROOT", "OPS", "ADMIN"))):
    """上传文档到知识库：自动解析、切片、嵌入、存入向量库"""
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
    if not kb:
        raise HTTPException(404, "知识库不存在")

    saved = await save_validated_upload(
        file,
        category=f"kb/{kb_id}",
        allowed_extensions=KB_EXTENSIONS,
        max_size=50 * 1024 * 1024,
    )
    ext = saved.extension
    file_path = str(saved.absolute_path)

    # 解析文件内容
    text = _parse_file(file_path, ext)

    # 创建文档记录
    doc = KBDocument(
        kb_id=kb_id, filename=saved.original_name,
        file_type=ext.replace(".", ""), file_size=saved.size,
        content=text[:50000], chunk_count=0, status="processing",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # 异步处理嵌入
    try:
        await _embed_document(kb, doc, text, db)
    except Exception as e:
        doc.status = "error"
        doc.error_msg = str(e)[:500]
        db.commit()
        raise HTTPException(500, f"文档处理失败: {e}")

    kb.doc_count = (kb.doc_count or 0) + 1
    db.commit()
    return {"id": doc.id, "filename": doc.filename, "chunk_count": doc.chunk_count, "status": doc.status}


@kb_router.delete("/{kb_id}/docs/{doc_id}")
def delete_doc(kb_id: int, doc_id: int, db: SessionLocal = Depends(get_db),
               user: User = Depends(require_role("ROOT", "OPS", "ADMIN"))):
    doc = db.query(KBDocument).filter(KBDocument.id == doc_id, KBDocument.kb_id == kb_id).first()
    if not doc:
        raise HTTPException(404, "文档不存在")
    delete_doc_chunks(kb_id, doc_id)
    db.delete(doc)
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
    if kb:
        kb.doc_count = max(0, (kb.doc_count or 1) - 1)
    db.commit()
    return {"deleted": True}


# ============ 检索问答 ============
@kb_router.post("/search")
async def search_kb(body: KBQueryIn, db: SessionLocal = Depends(get_db),
                    user: User = Depends(get_current_user)):
    """知识库检索"""
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == body.kb_id).first()
    if not kb:
        raise HTTPException(404, "知识库不存在")

    # 获取嵌入模型
    if not kb.embedding_model_id:
        raise HTTPException(400, "知识库未配置嵌入模型")
    emb_model = get_model(kb.embedding_model_id, db)
    if not emb_model or not emb_model.is_active:
        raise HTTPException(400, "嵌入模型不可用")

    # 查询向量化
    if "placeholder" in (emb_model.api_key or "") or len(emb_model.api_key or "") < 20:
        return {"question": body.question, "chunks": [],
                "answer": "嵌入模型 API Key 为占位符，请先在模型管理中配置真实的 embedding 模型 API Key"}
    vectors = await get_embedding(
        [body.question], emb_model.api_key, emb_model.endpoint, emb_model.model_name
    )
    if not vectors:
        return {"question": body.question, "chunks": [],
                "answer": f"嵌入服务调用失败（API返回异常），请检查嵌入模型 [{emb_model.name}] 的配置是否正确"}

    # Milvus 检索
    chunks = search_chunks(body.kb_id, vectors[0], top_k=body.top_k * 2 if body.use_rerank else body.top_k)
    if not chunks:
        return {"question": body.question, "chunks": [], "answer": "未找到相关内容"}

    # 重排序（可选）
    if body.use_rerank and kb.rerank_model_id:
        rerank_model = get_model(kb.rerank_model_id, db)
        if rerank_model and rerank_model.is_active:
            docs_text = [c["chunk_text"] for c in chunks]
            reranked = await rerank_documents(body.question, docs_text, rerank_model.api_key,
                                              rerank_model.endpoint, rerank_model.model_name, body.top_k)
            # 按 rerank 结果重新排序
            idx_map = {r["index"]: r["score"] for r in reranked}
            chunks.sort(key=lambda c: idx_map.get(chunks.index(c), 0), reverse=True)
            chunks = chunks[:body.top_k]

    return {
        "question": body.question,
        "chunks": [{"text": c["chunk_text"], "score": c.get("score", 0),
                     "doc_id": c.get("doc_id")} for c in chunks],
    }


@kb_router.post("/chat")
async def kb_chat(body: KBQueryIn, db: SessionLocal = Depends(get_db),
                  user: User = Depends(get_current_user)):
    """RAG 问答：检索 + LLM 生成回答"""
    from core.openai_client import OpenAIClient
    from dao.model_dao import get_default_model

    # 检索（search_kb 现在优雅降级，不再抛 500）
    try:
        search_result = await search_kb(body, db, user)
    except HTTPException:
        search_result = {"chunks": [], "answer": "检索服务暂不可用"}
    except Exception:
        search_result = {"chunks": [], "answer": "检索失败，请检查嵌入模型配置"}
    chunks = search_result.get("chunks", []) if isinstance(search_result, dict) else []
    context = "\n\n---\n\n".join([c["text"] for c in chunks]) if chunks else ""

    # 获取默认 chat 模型生成回答
    chat_model = get_default_model(db) or db.query(AIModel).filter(
        AIModel.model_type == "chat", AIModel.is_active == True
    ).first()

    if not chat_model:
        return {"question": body.question, "chunks": chunks, "answer": context or "无相关内容和可用模型"}

    if "placeholder" in chat_model.api_key or len(chat_model.api_key) < 20:
        return {"question": body.question, "chunks": chunks,
                "answer": _mock_rag_answer(body.question, chunks)}

    client = OpenAIClient(api_key=chat_model.api_key, endpoint=chat_model.endpoint,
                          model_name=chat_model.model_name, temperature=0.3, max_tokens=1024, timeout=30)
    try:
        messages = [
            {"role": "system", "content": "你是一个知识库助手。请根据提供的上下文回答用户问题。如果上下文中没有相关信息，请诚实告知。"},
            {"role": "user", "content": f"上下文：\n{context}\n\n问题：{body.question}"},
        ]
        resp = await client.chat_completion(messages)
        answer = OpenAIClient.extract_content(resp)
    except Exception:
        answer = _mock_rag_answer(body.question, chunks)
    finally:
        await client.close()

    return {"question": body.question, "chunks": chunks, "answer": answer}


# ============ 辅助函数 ============
def _parse_file(filepath: str, ext: str) -> str:
    """解析文件为纯文本"""
    try:
        if ext in {".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".py", ".js", ".ts"}:
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        elif ext == ".pdf":
            try:
                import fitz  # pymupdf
                doc = fitz.open(filepath)
                text = "\n".join([page.get_text() for page in doc])
                doc.close()
                return text
            except ImportError:
                return "[PDF 解析需要安装 pymupdf]"
        elif ext in {".docx", ".doc"}:
            try:
                from docx import Document
                doc = Document(filepath)
                return "\n".join([p.text for p in doc.paragraphs])
            except ImportError:
                return "[DOCX 解析需要安装 python-docx]"
        else:
            return f"[不支持的文件类型: {ext}]"
    except Exception as e:
        return f"[文件解析失败: {e}]"


async def _embed_document(kb: KnowledgeBase, doc: KBDocument, text: str, db):
    """对文档文本进行切片、嵌入并存入 Milvus"""
    # 切片
    chunks = chunk_text(text, kb.chunk_size, kb.chunk_overlap)
    if not chunks:
        doc.status = "error"
        doc.error_msg = "无法切分文本"
        db.commit()
        return

    # 获取嵌入模型
    if not kb.embedding_model_id:
        doc.status = "error"
        doc.error_msg = "未配置嵌入模型"
        db.commit()
        return
    emb_model = get_model(kb.embedding_model_id, db)
    if not emb_model:
        doc.status = "error"
        doc.error_msg = "嵌入模型不存在"
        db.commit()
        return

    # 批量嵌入
    chunk_texts = [c["text"] for c in chunks]
    vectors = await get_embedding(chunk_texts, emb_model.api_key, emb_model.endpoint, emb_model.model_name)
    if not vectors:
        doc.status = "error"
        doc.error_msg = "嵌入服务调用失败"
        db.commit()
        return

    # 存入 Milvus
    insert_chunks(kb.id, doc.id, chunks, vectors)

    doc.chunk_count = len(chunks)
    doc.status = "done"
    db.commit()


def _mock_rag_answer(question: str, chunks: list) -> str:
    if not chunks:
        return "知识库中暂无相关信息，请先上传文档。"
    preview = "\n".join([f"- {c['text'][:100]}..." for c in chunks[:3]])
    return f"根据知识库内容，与您的问题相关的信息如下：\n\n{preview}\n\n当前回答基于本地检索结果生成。"
